import io
import json
import re
import threading
from collections import OrderedDict
from pathlib import Path

import fitz  # PyMuPDF
import docx
import openpyxl
from bs4 import BeautifulSoup

try:
    from langchain_text_splitters import RecursiveCharacterTextSplitter
except ImportError:
    from langchain_text_splitters import RecursiveCharacterTextSplitter

import chromadb
from chromadb.utils.embedding_functions import SentenceTransformerEmbeddingFunction
import groq

# ── 定数 ─────────────────────────────────────────────────────────────────────
SIMILARITY_THRESHOLD = 0.8   # これ以上なら「見つからない」判定
MAX_SEARCH_RESULTS = 8       # ChromaDB から取得する上位件数（ベクトルスコア順）
MAX_UPLOAD_FILES = 5         # アップロード上限
MAX_FILE_SIZE_MB = 10        # ファイルサイズ上限
QUERY_EXPAND_THRESHOLD = 10  # 文字数がこれ以下なら expand_query() 実行
HISTORY_TURNS = 1            # Groqに渡す直前ターン数
GROQ_MODEL = "llama-3.3-70b-versatile"       # 回答生成（品質重視）
GROQ_MODEL_LIGHT = "llama-3.1-8b-instant"   # リランキング・クエリ補完（TPD 1,000,000/day）
WIKI_DIR = Path("data/wiki")
EMBED_MODEL = "paraphrase-multilingual-MiniLM-L12-v2"

# ── 埋め込み関数（初回ロード時にキャッシュ） ─────────────────────────────────
_embedding_fn = None


def get_embedding_fn():
    global _embedding_fn
    if _embedding_fn is None:
        _embedding_fn = SentenceTransformerEmbeddingFunction(model_name=EMBED_MODEL)
    return _embedding_fn


# ── Wiki 用プロセス共有コレクション（全セッションで1つだけ） ──────────────────
_wiki_client = None
_wiki_collection = None
_wiki_lock = threading.Lock()


def get_wiki_collection():
    """プロセス全体で1つだけ保持するWikiコレクション（スレッドセーフ初期化）"""
    global _wiki_client, _wiki_collection
    with _wiki_lock:
        if _wiki_client is None:
            _wiki_client = chromadb.Client()
            _wiki_collection = _wiki_client.get_or_create_collection(
                name="shisho_wiki",
                embedding_function=get_embedding_fn(),
                metadata={"hnsw:space": "cosine"},
            )
    return _wiki_collection


# ── セッション固有のアップロードコレクション ──────────────────────────────────
def init_upload_collection(session_id: str):
    """セッションごとに独立したアップロード用コレクション"""
    client = chromadb.Client()
    collection = client.get_or_create_collection(
        name=f"shisho_upload_{session_id}",
        embedding_function=get_embedding_fn(),
        metadata={"hnsw:space": "cosine"},
    )
    return client, collection


# ── テキスト分割設定 ──────────────────────────────────────────────────────────
def _get_splitter():
    return RecursiveCharacterTextSplitter(
        chunk_size=500,
        chunk_overlap=50,
        separators=["\n\n", "\n", "。", "、", ""],
    )


# ── ファイルのバイト読み込みヘルパー ─────────────────────────────────────────
def _read_bytes(file):
    if hasattr(file, "read"):
        return file.read()
    with open(file, "rb") as f:
        return f.read()


# ── テキスト抽出（ファイル種別で分岐） ──────────────────────────────────────
def extract_text(file, file_type, source_name=None):
    """
    ファイルからテキストを抽出してチャンクのリストを返す。
    各チャンク: {"text": str, "metadata": dict}

    file_type: "pdf" / "docx" / "xlsx" / "html"
    # v2予定: "url"（未実装）
    """
    if file_type == "pdf":
        return _extract_pdf(file, source_name)
    elif file_type == "docx":
        return _extract_docx(file, source_name)
    elif file_type == "xlsx":
        return _extract_xlsx(file, source_name)
    elif file_type == "html":
        return _extract_html(file, source_name)
    # elif file_type == "url": pass  # v2で実装
    return []


def _extract_pdf(file, source_name):
    chunks = []
    try:
        data = _read_bytes(file)
        doc = fitz.open(stream=data, filetype="pdf")
        page_texts = [(i + 1, page.get_text()) for i, page in enumerate(doc)]
        all_text = "".join(t for _, t in page_texts)

        # スキャンPDF判定
        if len(all_text.strip()) < 100:
            return [{"text": "__SCAN_PDF__", "metadata": {
                "source": source_name or "unknown.pdf", "file_type": "pdf"
            }}]

        splitter = _get_splitter()
        for page_num, text in page_texts:
            if not text.strip():
                continue
            for chunk in splitter.split_text(text):
                if chunk.strip():
                    chunks.append({"text": chunk, "metadata": {
                        "source": source_name or "unknown.pdf",
                        "file_type": "pdf",
                        "page": page_num,
                    }})
    except Exception:
        pass
    return chunks


def _extract_docx(file, source_name):
    chunks = []
    try:
        data = _read_bytes(file)
        document = docx.Document(io.BytesIO(data))
        paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
        full_text = "\n".join(paragraphs)

        splitter = _get_splitter()
        for i, chunk in enumerate(splitter.split_text(full_text)):
            if chunk.strip():
                chunks.append({"text": chunk, "metadata": {
                    "source": source_name or "unknown.docx",
                    "file_type": "docx",
                    "page": i + 1,
                }})
    except Exception:
        pass
    return chunks


def _extract_xlsx(file, source_name):
    chunks = []
    try:
        data = _read_bytes(file)
        wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = []
            for row in ws.iter_rows(values_only=True):
                row_text = [str(cell) if cell is not None else "" for cell in row]
                if any(c.strip() for c in row_text):
                    rows.append(row_text)
            if not rows:
                continue

            header = rows[0]
            data_rows = rows[1:]
            chunk_rows = 10

            for start in range(0, max(1, len(data_rows)), chunk_rows):
                end = min(start + chunk_rows, len(data_rows))
                lines = [
                    f"シート: {sheet_name}",
                    "ヘッダー: " + " | ".join(str(h) for h in header if str(h).strip()),
                ]
                lines += [" | ".join(str(c) for c in row) for row in data_rows[start:end]]
                text = "\n".join(lines)
                if text.strip():
                    chunks.append({"text": text, "metadata": {
                        "source": source_name or "unknown.xlsx",
                        "file_type": "xlsx",
                        "sheet": sheet_name,
                        "rows": f"{start + 2}-{end + 1}",
                    }})
    except Exception:
        pass
    return chunks


# ── HTML ノイズ除去定数（RAG 精度向上） ─────────────────────────────────────
HTML_NOISE_CLASSES = [
    "ornamental-border",   # — ✦ ——— ✦ — などの飾り罫
    "title-divider",       # タイトル下の区切り線
    "chapter-label",       # "Catalogue of Works" などの英語ラベル
    "page-footer",         # "p.3 / Yuruyuru Series — Scholar's Study / 2026"
    "section-divider",     # · · · などの区切り
    "footer-series-name",  # フッターのシリーズ名
    "page-number",         # フッターのページ番号
    "app-url-link",        # "▶ デモを開く" リンク
    "tech-badge",          # 技術スタックバッジ（Streamlit, Groq等）
    "chara-badge",         # "主役" "宛先" などのバッジ
]
HTML_NOISE_TAGS = ["style", "script", "head"]


def _extract_html(file, source_name):
    chunks = []
    try:
        data = _read_bytes(file)
        content = data.decode("utf-8", errors="replace") if isinstance(data, bytes) else data
        soup = BeautifulSoup(content, "html.parser")

        # タイトル・見出しをノイズ除去前に取得
        title_tag = soup.find("title")
        title = title_tag.get_text(strip=True) if title_tag else ""
        headings = [tag.get_text(strip=True) for tag in soup.find_all(["h1", "h2"]) if tag.get_text(strip=True)]
        heading = " / ".join(headings)

        # ノイズ除去（装飾・ナビゲーション・スタイル系）
        for cls in HTML_NOISE_CLASSES:
            for tag in soup.find_all(class_=cls):
                tag.decompose()
        for tag_name in HTML_NOISE_TAGS:
            for tag in soup.find_all(tag_name):
                tag.decompose()

        text = re.sub(r"\n{3,}", "\n\n", soup.get_text(separator="\n", strip=True))

        # 100文字未満はスキップ（装飾のみのページ対策）
        if len(text.strip()) < 100:
            return chunks

        meta = {
            "source": source_name or "unknown.html",
            "file_type": "html",
            "title": title,
            "heading": heading,
        }

        # 500文字以下は分割しない
        if len(text) <= 500:
            chunks.append({"text": text, "metadata": meta})
        else:
            for chunk in _get_splitter().split_text(text):
                if chunk.strip() and len(chunk.strip()) >= 100:
                    chunks.append({"text": chunk, "metadata": meta})
    except Exception:
        pass
    return chunks


# ── ChromaDB への登録 ─────────────────────────────────────────────────────────
def store_chunks(collection, chunks, source_type="wiki"):
    if not chunks:
        return

    documents, metadatas, ids = [], [], []
    for i, chunk in enumerate(chunks):
        if chunk["text"] == "__SCAN_PDF__":
            continue
        meta = {
            k: (v if isinstance(v, (str, int, float, bool)) else str(v))
            for k, v in chunk["metadata"].items()
        }
        meta["source_type"] = source_type
        doc_id = f"{source_type}_{meta.get('source', 'unknown')}_{i}"
        documents.append(chunk["text"])
        metadatas.append(meta)
        ids.append(doc_id)

    if not documents:
        return

    batch_size = 100
    for start in range(0, len(documents), batch_size):
        collection.upsert(
            documents=documents[start:start + batch_size],
            metadatas=metadatas[start:start + batch_size],
            ids=ids[start:start + batch_size],
        )


# ── Wiki 読み込み ─────────────────────────────────────────────────────────────
EXT_MAP = {".pdf": "pdf", ".docx": "docx", ".xlsx": "xlsx", ".html": "html", ".htm": "html"}


def load_wiki_to_chromadb(wiki_collection, progress_callback=None):
    """WIKI_DIR の全ファイルを wiki_collection に登録する。
    既にデータがある場合（別セッションがロード済み）はスキップする。
    """
    if wiki_collection.count() > 0:
        return  # 別セッションが既にロード済み

    if not WIKI_DIR.exists():
        return

    files = [
        f for f in sorted(WIKI_DIR.iterdir())
        if f.is_file() and not f.name.startswith(".") and f.suffix.lower() in EXT_MAP
    ]
    total = len(files)

    for done, filepath in enumerate(files, 1):
        file_type = EXT_MAP[filepath.suffix.lower()]
        chunks = extract_text(filepath, file_type, source_name=filepath.name)
        store_chunks(wiki_collection, chunks, source_type="wiki")
        if progress_callback:
            progress_callback(done, total, filepath.name)


# ── クエリ補完 ────────────────────────────────────────────────────────────────
def expand_query(query, groq_client):
    """10文字以下の短い質問のみ Groq で補完"""
    if len(query) > QUERY_EXPAND_THRESHOLD:
        return query
    try:
        prompt = (
            f'以下の短い質問を検索しやすい質問文に補完してください。JSONのみ返してください（前置き不要）。\n'
            f'質問: {query}\n{{"expanded": "補完した質問文"}}'
        )
        resp = groq_client.chat.completions.create(
            model=GROQ_MODEL_LIGHT,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3,
            max_tokens=100,
        )
        m = re.search(r"\{.*\}", resp.choices[0].message.content, re.DOTALL)
        if m:
            return json.loads(m.group()).get("expanded", query)
    except groq.RateLimitError:
        print(f"[Groq] expand_query: RateLimitError")
        raise
    except Exception as e:
        print(f"[Groq] expand_query: {type(e).__name__}: {e}")
    return query


# ── 会話要約（同期） ──────────────────────────────────────────────────────────
def summarize_conversation_sync(last_exchange, current_summary, groq_client):
    """直前の Q&A を要約して返す（次の質問処理の冒頭で呼ぶ）"""
    try:
        q = last_exchange.get("q", "")
        a = last_exchange.get("a", "")
        prev = current_summary or ""
        prompt = (
            f'会話を1〜2文で要約してください。JSONのみ返してください。\n'
            f'【これまでの要約】{prev}\n'
            f'【直前のやりとり】Q: {q}\nA: {a}\n'
            f'{{"summary": "1〜2文の要約"}}'
        )
        resp = groq_client.chat.completions.create(
            model=GROQ_MODEL,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3,
            max_tokens=150,
        )
        m = re.search(r"\{.*\}", resp.choices[0].message.content, re.DOTALL)
        if m:
            return json.loads(m.group()).get("summary", current_summary)
    except Exception:
        pass
    return current_summary


# ── 類似検索 ──────────────────────────────────────────────────────────────────
def _query_collection(collection, query):
    """1つのコレクションを検索してヒットリストを返す（内部ヘルパー）"""
    try:
        count = collection.count()
        if count == 0:
            return []
        results = collection.query(
            query_texts=[query],
            n_results=min(MAX_SEARCH_RESULTS, count),
        )
        if not results or not results["documents"] or not results["documents"][0]:
            return []
        return list(zip(
            results["documents"][0],
            results["metadatas"][0],
            results["distances"][0],
        ))
    except Exception:
        return []


def search_documents(wiki_collection, upload_collection, query):
    """
    WikiコレクションとUploadコレクションを両方検索してマージする。
    見つからない場合は None を返す。
    """
    wiki_hits = _query_collection(wiki_collection, query)
    upload_hits = _query_collection(upload_collection, query) if upload_collection else []

    all_hits = wiki_hits + upload_hits
    if not all_hits:
        return None

    # 距離の昇順にソート
    all_hits.sort(key=lambda x: x[2])

    # 最小距離が閾値以上なら「見つからない」
    if all_hits[0][2] >= SIMILARITY_THRESHOLD:
        return None

    # 重複排除（同一ファイル＋同一スコア）して上位 MAX_SEARCH_RESULTS 件を返す
    seen_keys = set()
    final = []
    for doc, meta, dist in all_hits:
        key = (meta.get("source", ""), round(dist, 6))
        if key not in seen_keys:
            seen_keys.add(key)
            final.append({"text": doc, "metadata": meta, "distance": dist})
        if len(final) >= MAX_SEARCH_RESULTS:
            break

    # デバッグ：取得件数と距離一覧を出力
    print(f"[Search] {len(final)} hits:")
    for i, item in enumerate(final):
        src = item["metadata"].get("source", "不明")
        print(f"  [{i}] dist={item['distance']:.4f}  {src}")

    return final or None


# ── 回答生成 ──────────────────────────────────────────────────────────────────
def generate_response(query, search_results, chat_summary, groq_client):
    """Groq にホーさん口調のまとめ＋ハイライトキーワードを生成させる。
    戻り値: (intro: str, keywords: list[str])
    """
    context = "\n\n".join(
        f"[{r['metadata'].get('source', '不明')}]\n{r['text']}" for r in search_results
    )
    summary_text = f"\n【これまでの会話の要約】\n{chat_summary}" if chat_summary else ""
    prompt = (
        f"あなたはホーさん🦉（フクロウの司書）です。"
        f"「〜だねぇ」「えーとえーと」「あったあった〜」のようなゆったり口調で話します。"
        f"{summary_text}\n\n"
        f"以下の検索結果を読んで、JSONのみ返してください（前置き・マークダウン不要）。\n\n"
        f"【質問】{query}\n"
        f"【検索結果】\n{context}\n\n"
        f"【条件】\n"
        f"・質問への答えがチャンクに含まれている場合は全て具体的に答えること\n"
        f"・渡されたチャンクに書かれていることだけで答え、知っていても書かれていないことは答えないこと\n"
        f"・keywords は【質問への答え】に直接関係する単語で最大5個。存在しない単語は入れないこと\n\n"
        f'{{"intro": "ホーさんのコメント（80字以内、ゆったり口調）", '
        f'"keywords": ["質問の答えに関連するキーワード"]}}'
    )
    try:
        resp = groq_client.chat.completions.create(
            model=GROQ_MODEL,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7,
            max_tokens=200,
        )
        m = re.search(r"\{.*\}", resp.choices[0].message.content, re.DOTALL)
        if m:
            data = json.loads(m.group())
            intro = data.get("intro", "えーとえーと…見つかったよ〜🦉")
            keywords = [k for k in data.get("keywords", []) if isinstance(k, str) and k]
            return intro, keywords
    except groq.RateLimitError:
        print(f"[Groq] generate_response: RateLimitError")
        raise
    except Exception as e:
        print(f"[Groq] generate_response: {type(e).__name__}: {e}")
    return "えーとえーと…見つかったよ〜🦉", []


# ── 出典フォーマット ──────────────────────────────────────────────────────────
def format_source_item(item):
    """チャンク 1 件の表示用情報を dict で返す"""
    meta = item["metadata"]
    ft = meta.get("file_type", "")
    src = meta.get("source", "不明")

    if ft == "html":
        label = meta.get("title", "") or src
        return {"icon": "🌐", "label": label, "detail": None,
                "source": src, "file_type": ft, "text": item["text"]}
    elif ft == "pdf":
        pg = meta.get("page", "")
        return {"icon": "📄", "label": src, "detail": f"{pg}ページ目" if pg else "",
                "source": src, "file_type": ft, "text": item["text"]}
    elif ft == "docx":
        pg = meta.get("page", "")
        return {"icon": "📝", "label": src, "detail": f"{pg}ページ目" if pg else "",
                "source": src, "file_type": ft, "text": item["text"]}
    elif ft == "xlsx":
        sheet = meta.get("sheet", "")
        rows = meta.get("rows", "")
        detail = f"{sheet} ／ {rows}行目" if sheet and rows else sheet or rows
        return {"icon": "📊", "label": src, "detail": detail,
                "source": src, "file_type": ft, "text": item["text"]}
    return {"icon": "📎", "label": src, "detail": None,
            "source": src, "file_type": ft, "text": item["text"]}


def group_results_by_file(search_results):
    """検索結果をファイル名単位でグループ化して OrderedDict で返す"""
    groups = OrderedDict()
    for item in search_results:
        src = item["metadata"].get("source", "不明")
        if src not in groups:
            groups[src] = []
        groups[src].append(item)
    return groups


# ── 資料室用：Wiki ファイル情報 ───────────────────────────────────────────────
def get_wiki_file_info():
    """資料室タブ用に wiki ファイルの情報リストを返す"""
    if not WIKI_DIR.exists():
        return []

    files_info = []
    for f in sorted(WIKI_DIR.iterdir()):
        if not f.is_file() or f.name.startswith("."):
            continue
        ft = EXT_MAP.get(f.suffix.lower())
        if not ft:
            continue

        info = {"name": f.name, "file_type": ft, "path": str(f)}

        if ft == "html":
            try:
                content = f.read_text(encoding="utf-8", errors="replace")
                soup = BeautifulSoup(content, "html.parser")
                t = soup.find("title")
                info["title"] = t.get_text(strip=True) if t else f.stem
            except Exception:
                info["title"] = f.stem

        elif ft == "pdf":
            try:
                doc = fitz.open(str(f))
                info["pages"] = doc.page_count
            except Exception:
                info["pages"] = None

        elif ft == "docx":
            try:
                document = docx.Document(str(f))
                paras = len([p for p in document.paragraphs if p.text.strip()])
                info["pages"] = max(1, paras // 10)
            except Exception:
                info["pages"] = None

        files_info.append(info)

    return files_info
